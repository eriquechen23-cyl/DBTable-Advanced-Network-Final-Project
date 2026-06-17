from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results" / "ta_dbtable_results.json"
DELIVERABLES = ROOT / "deliverables"
REPORTS = ROOT / "docs" / "reports"
STUDENT = "P77141155 陳燁龍"


def normalize_summary(raw: dict) -> dict:
    summary = dict(raw)
    summary["build_seconds"] = raw.get("build_seconds_avg", raw.get("build_seconds", 0.0))
    summary["build_seconds_min"] = raw.get("build_seconds_min", summary["build_seconds"])
    summary["lookup_avg_ns"] = raw.get("lookup_avg_ns_avg", raw.get("lookup_avg_ns", 0.0))
    summary["lookup_avg_ns_min"] = raw.get("lookup_avg_ns_min", summary["lookup_avg_ns"])
    summary["lookup_total_seconds"] = raw.get(
        "lookup_total_seconds_avg", raw.get("lookup_total_seconds", 0.0)
    )
    summary["estimated_memory_mib"] = raw.get(
        "estimated_memory_mib", raw.get("estimated_memory_bytes", 0) / 1024 / 1024
    )
    return summary


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "Microsoft JhengHei"
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_para(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix and text.startswith(bold_prefix):
        run = p.add_run(bold_prefix)
        run.bold = True
        p.add_run(text[len(bold_prefix) :])
    else:
        p.add_run(text)


def set_cell(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_after = Pt(2)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell(table.rows[0].cells[i], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell(cells[i], value)
    doc.add_paragraph()


def build_docx(summary: dict) -> Path:
    DELIVERABLES.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    doc.styles["Normal"].font.name = "Microsoft JhengHei"
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("高效能網路期末專案書面報告\nDBTable Packet Classification")
    run.bold = True
    run.font.name = "Microsoft JhengHei"
    run.font.size = Pt(22)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"演算法：DBTable | 資料集：ClassBench ACL1 100K rules | 作者：{STUDENT}")
    doc.add_paragraph()

    add_heading(doc, "摘要", 1)
    add_para(
        doc,
        "本報告研究 DBTable 封包分類演算法。封包分類的目標，是在大量規則中依照來源 IP、目的 IP、來源埠、目的埠與協定，找出最高優先權的相符規則。DBTable 的重點不是把每條規則都拿來逐一檢查，而是先從規則集合中找出最有辨識力的幾個 bit，利用這些 bit 把規則分到多個小表格中。查詢封包時，只要先查到對應的小表格，再做完整五元組驗證，就能減少平均查找成本。",
    )

    add_heading(doc, "1. 論文方法的白話說明", 1)
    add_heading(doc, "1.1 一分鐘白話版", 2)
    add_para(
        doc,
        "可以把 DBTable 想成一個大型圖書館的快速索引。沒有索引時，找一本書要從第一排一路掃到最後一排；DBTable 的做法是先觀察所有書的特徵，挑出最能分辨書本位置的幾個欄位，做成索引卡。封包來時，先用這張索引卡找到可能的位置，再到那一小格裡逐一確認。",
    )
    add_para(
        doc,
        "套回封包分類：書本就是 rules，索引卡就是 discriminative bitsets，小格子就是 bucket。DBTable 希望每個 bucket 裡的候選規則越少越好，但又不能漏掉真正會匹配的規則。因此它會讓有 wildcard 的規則複製到多個可能 bucket，查詢時再做精確比對，保證結果仍是最高優先權的正確規則。",
    )

    add_heading(doc, "1.2 DBTable 實際做的四件事", 2)
    add_table(
        doc,
        ["步驟", "白話意思", "技術意義"],
        [
            ["1. 觀察 ruleset", "先看所有規則的 IP prefix 分布，不急著查詢。", "統計來源/目的 IP bit 在規則中的分布與可分辨程度。"],
            ["2. 挑出有辨識力的 bit", "挑幾個最能把規則分開的問題，例如第 1 個 bit 是 0 還是 1。", "形成 discriminative bitset，讓 bucket 盡量平均、候選集合盡量小。"],
            ["3. 建立 bucket table", "每條規則依照這些 bit 被放到對應格子；如果規則太寬，就放到多個格子。", "處理 prefix wildcard 與 replication，建立查詢用索引表。"],
            ["4. 查詢時先定位再驗證", "封包先用相同 bit 找格子，再只檢查格子裡的規則。", "candidate reduction 後仍做 exact five-tuple match 與 priority selection。"],
        ],
    )

    add_heading(doc, "1.3 為什麼這樣會快", 2)
    add_para(
        doc,
        "封包分類慢，通常不是因為判斷一條規則很難，而是因為規則太多。若 ruleset 有十萬筆，每個封包都掃十萬筆會很慢。DBTable 的加速點在於：先用少數 bit 做快速索引，把候選規則從十萬筆縮到一個 bucket 裡的幾百筆，最後才做完整五元組比對。",
    )
    add_para(
        doc,
        "這也解釋了 DBTable 的取捨：它用較多建表時間與額外記憶體，換取查詢時少掃很多規則。若規則的 IP prefix 分布很有規律，選到的 discriminative bits 可以把規則切得很開，DBTable 就會有好效果；若規則充滿 wildcard，規則會被複製到很多 bucket，記憶體與建表成本就會上升。",
    )

    add_heading(doc, "1.4 與論文名詞對照", 2)
    add_table(
        doc,
        ["論文/程式名詞", "白話解釋", "在本專案中的對應"],
        [
            ["Discriminative bitsets", "最能把規則分群的 bit 集合。", "src/dbtable_classifier.py 會對 64 個來源/目的 IP bit 排名並選出 12 個。"],
            ["Bucket table", "查詢用索引表，每個 key 對到一小群候選規則。", "buckets: dict[int, list[int]]。"],
            ["Wildcard / prefix replication", "規則沒有指定某個 bit 時，它可能屬於多個 bucket。", "_rule_signatures() 會產生多個相容 key。"],
            ["Exact match", "最後仍要完整檢查五元組，不能只靠索引猜答案。", "rule_matches() 檢查 IP prefix、port range、protocol mask。"],
            ["Highest priority rule", "多條規則匹配時，回傳優先權最高者。", "bucket 內依 priority 排序，找到最早匹配規則。"],
        ],
    )

    add_heading(doc, "2. 演算法介紹與架構分析", 1)
    add_heading(doc, "2.1 問題定義", 2)
    add_para(
        doc,
        "封包分類輸入為 packet 五元組：source IP、destination IP、source port、destination port、protocol。規則包含 IP prefix、port range、protocol mask 與 priority。查詢結果必須回傳所有相符規則中優先權最高者。",
    )

    add_heading(doc, "2.2 架構流程", 2)
    diagram = doc.add_paragraph()
    diagram.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = diagram.add_run(
        "ClassBench Rules -> Parser -> Discriminative Bit Ranking -> Bucket Table Build\n"
        "Packet Trace -> Key Extraction -> Candidate Bucket -> Exact 5-tuple Match -> Highest Priority Rule"
    )
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    add_table(
        doc,
        ["階段", "輸入", "主要工作", "輸出"],
        [
            ["Rule parsing", "ClassBench rule file", "解析 IP prefix、port range、protocol mask、priority。", "Rule objects"],
            ["Bit ranking", "All rules", "計算每個 IP bit 的分割平衡度與覆蓋率。", "Selected bits"],
            ["Bucket build", "Rules + selected bits", "依 selected bits 建立 bucket；wildcard bit 會複製到多個 bucket。", "Bucket table"],
            ["Lookup", "Packet trace", "取 packet key，掃描候選 bucket，做精確五元組比對。", "Best matching priority"],
        ],
    )

    add_heading(doc, "3. 實作方法", 1)
    add_para(
        doc,
        "本專案原先使用 Python 教學版說明 DBTable 概念；本次依照助教提供的 DBTable.cpp / AMPS 官方 src/DBTable.cpp 改為 C++ 實測。benchmark harness 位於 cpp/benchmark_ta_dbtable.cpp，負責解析 ClassBench rules/trace、建立 DBT::DBTable、量測 build time、lookup time 與 memory estimate。可重現腳本為 scripts/run_ta_dbtable_experiment.py。",
    )
    add_table(
        doc,
        ["檔案", "用途"],
        [
            ["external/amps/src/DBTable.cpp", "AMPS 官方 DBTable C++ 實作，與助教提供來源對應。"],
            ["cpp/benchmark_ta_dbtable.cpp", "C++ benchmark harness：讀取 ClassBench、建立 DBTable、量測效能。"],
            ["scripts/run_ta_dbtable_experiment.py", "自動編譯並重複執行 5 次，輸出 ta_dbtable_results.json/csv。"],
            ["ta_reference/DBTable.cpp", "保留助教給的原始檔作對照；本地複製檔有一行 Init 被註解吃掉，因此實測使用 AMPS 官方檔。"],
            ["src/dbtable_classifier.py", "早期 Python 教學版，保留作概念對照，不作正式數據來源。"],
        ],
    )

    add_heading(doc, "4. 實驗與效能評估", 1)
    add_heading(doc, "4.1 資料集與執行方式", 2)
    add_para(
        doc,
        f"使用 ClassBench ACL1 100K rules 檔案 data/classbench/acl1_100000.txt 與對應 trace。此下載檔名為 100K，實際有效可解析規則數為 {summary['rules_loaded']:,} 筆，trace 測試封包數為 {summary['packets_tested']:,} 筆。",
    )
    add_para(doc, "執行指令：python scripts/run_experiment.py --max-traces 100000")
    add_heading(doc, "4.2 實驗結果", 2)
    add_table(
        doc,
        ["指標", "結果"],
        [
            ["Implementation", summary.get("implementation", "AMPS/TA DBTable.cpp")],
            ["Repeat count", str(summary.get("repeats", 1))],
            ["DBTable threshold (BINTH)", str(summary.get("threshold", 4))],
            ["Build time avg", f"{summary['build_seconds']:.4f} s"],
            ["Build time min", f"{summary['build_seconds_min']:.4f} s"],
            ["Lookup total time avg", f"{summary['lookup_total_seconds']:.6f} s"],
            ["Average lookup time avg", f"{summary['lookup_avg_ns']:.3f} ns"],
            ["Average lookup time min", f"{summary['lookup_avg_ns_min']:.3f} ns"],
            ["Estimated memory", f"{summary['estimated_memory_mib']:.3f} MiB"],
        ],
    )
    add_para(
        doc,
        "結果顯示，改用 AMPS/助教 C++ DBTable.cpp 後，lookup time 大幅下降到數十 ns 等級。這代表先前 Python 版本主要適合說明概念，不適合作為正式效能比較。正式比較應以本節 C++ 實測數字為準。",
    )

    add_heading(doc, "5. 與組員的比較分析", 1)
    add_para(
        doc,
        "本節比較 DBTable、HybridTSS 與 CutSplit 三種方法。DBTable 數據已改用 AMPS/助教 C++ DBTable.cpp 重新量測，與先前 Python 教學版不同。三者都使用 ClassBench acl1_100k 類資料，但 CutSplit 回報實際載入 99,833 筆規則，DBTable 本次載入 99,330 筆規則，因此仍需注意資料與實作細節差異。",
    )
    add_table(
        doc,
        ["面向", "DBTable（本專案）", "HybridTSS（組員一）", "CutSplit（組員二）"],
        [
            ["設計核心", "用 discriminative bits 建 bucket，先縮小候選集合。", "以 tuple space/hash 類索引快速定位相同 tuple 的候選規則。", "以 cut/split 決策結構切分搜尋空間，讓封包沿路徑找到候選規則。"],
            ["Rules loaded", f"{summary['rules_loaded']:,}", "未提供", "99,833"],
            ["Build time", f"{summary['build_seconds']:.4f} s", "0.0249 s", "0.4625 s"],
            ["Average lookup time", f"{summary['lookup_avg_ns']:.3f} ns", "133.987 ns", "300.719 ns"],
            ["Memory", f"{summary['estimated_memory_mib']:.3f} MiB", "709.45 MiB", "0.512 MiB"],
            ["優勢", "lookup time 最低，memory 也遠低於 HybridTSS。", "build time 最短。", "memory 最低，build/lookup/memory 整體均衡。"],
            ["弱點", "build time 不是最低，memory 高於 CutSplit。", "memory consumption 很高，約為 DBTable 的 176.6 倍。", "lookup 比 DBTable 與 HybridTSS 慢。"],
            ["適用情境", "需要極低 lookup latency，同時希望 memory 不像 HybridTSS 那麼高。", "需要極快建表且記憶體充足。", "低記憶體設備或整體平衡需求。"],
        ],
    )
    add_para(
        doc,
        f"從更新後的 C++ 實測數字看，DBTable 平均 lookup time 為 {summary['lookup_avg_ns']:.3f} ns，是三者中最低；HybridTSS 為 133.987 ns，CutSplit 為 300.719 ns。HybridTSS 的 build time 只有 0.0249 s，是三者最快，但 memory 達 709.45 MiB，成本最高。CutSplit 的 memory 只有 0.512 MiB，是三者最低，適合記憶體受限情境。整體來看，DBTable 在本次比較中提供最佳查詢延遲與可接受的記憶體；CutSplit 則是最佳低記憶體選擇；HybridTSS 適合極快建表但記憶體充足的情境。",
    )

    add_heading(doc, "6. 結論", 1)
    add_para(
        doc,
        "DBTable 的價值在於把 ruleset 的分布特性轉換成快速 bucket index，使 packet lookup 不必面對完整規則集合。本次實驗在 ClassBench ACL1 100K 資料上完成建表、查詢與記憶體估計，並透過線性掃描抽樣驗證正確性。後續若要更貼近論文與 AMPS 實作，可將 prefix_tuple、port_node、動態 tuple range 與 C++ SIMD/hash optimization 納入。",
    )

    add_heading(doc, "參考資料", 1)
    for ref in [
        "DBTable: Leveraging Discriminative Bitsets for High-Performance Packet Classification, IEEE/ACM Transactions on Networking, 2024.",
        "Packet classification algorithms list: https://github.com/matthiola0/packet-classification-algorithms",
        "AMPS reference implementation including DBTable: https://github.com/JiaChangGit/amps",
        "ClassBench packet classification generator: external/classbench-packet-classification and https://github.com/classbench-ng/classbench-ng",
    ]:
        doc.add_paragraph(ref)

    out = DELIVERABLES / "DBTable_Final_Project_Report.docx"
    doc.save(out)
    return out


def build_html(summary: dict) -> Path:
    REPORTS.mkdir(parents=True, exist_ok=True)
    html = f"""<!doctype html>
<html lang="zh-Hant">
<head>
  <meta charset="utf-8">
  <title>DBTable Architecture Milestone Report</title>
  <style>
    body {{ font-family: Arial, "Microsoft JhengHei", sans-serif; margin: 40px; color: #1f2937; line-height: 1.55; }}
    h1, h2 {{ color: #143d59; }}
    .box {{ border: 1px solid #9eb6c8; background: #f7fbff; border-radius: 8px; padding: 12px; margin: 8px 0; }}
    table {{ border-collapse: collapse; width: 100%; margin: 16px 0; }}
    th, td {{ border: 1px solid #d0d7de; padding: 8px; vertical-align: top; }}
    th {{ background: #eaf3f8; }}
    code {{ background: #f4f4f4; padding: 2px 4px; }}
  </style>
</head>
<body>
  <h1>DBTable Architecture Milestone Report</h1>
  <p>Project: Final-Network | Author: {STUDENT} | Milestone: AMPS/TA DBTable.cpp benchmark verified</p>
  <h2>Plain-Language Method</h2>
  <div class="box">DBTable is like a library index: choose a few highly discriminative IP bits, use them to jump to a small bucket, then perform exact five-tuple verification only inside that bucket.</div>
  <h2>Architecture Flow</h2>
  <p><code>Rules -> Parser -> Bit Ranking -> Bucket Table</code></p>
  <p><code>Packet -> Key Extraction -> Candidate Bucket -> Exact Match -> Highest Priority Rule</code></p>
  <h2>Benchmark Snapshot</h2>
  <table>
    <tr><th>Metric</th><th>Value</th></tr>
    <tr><td>Rules</td><td>{summary['rules_loaded']:,}</td></tr>
    <tr><td>Build time avg</td><td>{summary['build_seconds']:.4f} s</td></tr>
    <tr><td>Average lookup avg</td><td>{summary['lookup_avg_ns']:.3f} ns</td></tr>
    <tr><td>Estimated memory</td><td>{summary['estimated_memory_mib']:.3f} MiB</td></tr>
  </table>
</body>
</html>
"""
    out = REPORTS / "20260611-DBTable-architecture-milestone.html"
    out.write_text(html, encoding="utf-8")
    return out


def build_script(summary: dict) -> Path:
    text = f"""# DBTable 20 分鐘投影片報告講稿

## 0:00-1:30 封面與研究目標
本次選擇 DBTable 作為 packet classification 演算法。我的重點不是只展示程式結果，而是說清楚論文方法：DBTable 如何用 discriminative bitsets 把十萬筆規則切成較小的候選集合。

## 1:30-4:00 Packet Classification 問題
封包分類要看五元組：來源 IP、目的 IP、來源埠、目的埠、協定。每條規則有 prefix、range、mask 與 priority。查詢時必須找最高優先權的匹配規則。

## 4:00-8:30 DBTable 白話版
DBTable 像圖書館索引。沒有索引時要一本一本找；DBTable 先從規則中挑出最能分辨位置的 bit，做成索引。封包進來時先用這些 bit 找到小 bucket，再在 bucket 中做完整比對。這樣不會漏掉正確答案，因為最後仍然做 exact match。

## 8:30-11:30 為什麼會快、代價是什麼
快的原因是候選規則變少。代價是建表需要先分析 ruleset，而且 wildcard prefix 可能讓規則被放到多個 bucket，增加記憶體。也就是用建表成本和記憶體換查詢速度。

## 11:30-14:00 程式實作
本次正式數據使用 AMPS/助教 C++ DBTable.cpp。cpp/benchmark_ta_dbtable.cpp 負責把 ClassBench rules/trace 轉成 DBT::Rule 和 DBT::Packet，然後建立 DBT::DBTable 並量測 build、lookup、memory。scripts/run_ta_dbtable_experiment.py 會自動編譯並跑 5 次。

## 14:00-17:00 實驗結果
有效規則數 {summary['rules_loaded']:,}，測試封包 {summary['packets_tested']:,}，重複 {summary.get('repeats', 1)} 次。Build time avg {summary['build_seconds']:.4f} 秒，average lookup avg {summary['lookup_avg_ns']:.3f} ns，記憶體約 {summary['estimated_memory_mib']:.3f} MiB。

## 17:00-19:00 比較分析
組員一 HybridTSS：build time 0.0249 秒、average lookup 133.987 ns、memory 709.45 MiB。組員二 CutSplit：build time 0.4625 秒、average lookup 300.719 ns、memory 0.512 MiB。更新後 DBTable C++ 實測 lookup 約 {summary['lookup_avg_ns']:.3f} ns，是三者中最快；memory 約 {summary['estimated_memory_mib']:.3f} MiB，遠低於 HybridTSS，但高於 CutSplit。因此 DBTable 適合追求極低查詢延遲且可接受數 MiB 記憶體的情境。

## 19:00-20:00 結論
DBTable 的核心是一句話：先用有辨識力的 bit 快速定位候選 bucket，再用完整五元組驗證正確答案。本專案完成資料集、程式、實驗、報告與 PPT。
"""
    out = DELIVERABLES / "DBTable_20min_Presentation_Script.md"
    out.write_text(text, encoding="utf-8")
    return out


def main() -> None:
    summary = normalize_summary(json.loads(RESULTS.read_text(encoding="utf-8")))
    docx = build_docx(summary)
    html = build_html(summary)
    script = build_script(summary)
    print(json.dumps({"docx": str(docx), "html": str(html), "script": str(script)}, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
